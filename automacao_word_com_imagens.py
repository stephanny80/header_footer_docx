import os
import copy
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree


class DocxLayoutSynchronizer:
    """
    Orquestra a sincronização de layout e conteúdo de um documento Word de template
    para um documento de destino.

    Esta classe representa o caso de uso principal (Application Service no DDD).
    """
    _NSMAP = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    def __init__(self, template_path: str, dest_path: str):
        print("Iniciando processo de sincronização de layout e conteúdo...")
        self.template_doc = Document(template_path)
        self.dest_doc = Document(dest_path)
        self.style_handler = StyleHandler(self.template_doc, self._NSMAP)

    def run(self):
        """Executa o processo de sincronização para todas as seções."""
        for i, dest_section in enumerate(self.dest_doc.sections):
            if i >= len(self.template_doc.sections):
                break  # Não há mais seções de template para copiar

            template_section = self.template_doc.sections[i]
            print(f"\nProcessando Seção {i}...")

            # Delega o processamento da seção para uma classe especialista
            section_processor = SectionProcessor(template_section, dest_section, self._NSMAP)
            section_processor.copy_properties()

            # Processa cada parte (cabeçalho, rodapé, etc.)
            self._process_parts(section_processor, template_section, dest_section)

        print("\nProcesso de sincronização concluído.")

    def _process_parts(self, section_processor: 'SectionProcessor', template_section, dest_section):
        """Processa a cópia de todas as partes relevantes da seção."""
        parts_to_copy = {
            'default_header': (template_section.header, dest_section.header),
            'default_footer': (template_section.footer, dest_section.footer),
        }

        if section_processor.is_first_page_different():
            print("  - 'Primeira Página Diferente' detectado. Processando partes específicas.")
            parts_to_copy.update({
                'first_page_header': (template_section.first_page_header, dest_section.first_page_header),
                'first_page_footer': (template_section.first_page_footer, dest_section.first_page_footer),
            })

        for part_name, (source_el, dest_el) in parts_to_copy.items():
            copier = PartCopier(source_el, dest_el, self.style_handler, self._NSMAP, part_name)
            copier.copy_content()

    def save(self, output_path: str):
        """Salva o documento de destino modificado."""
        try:
            self.dest_doc.save(output_path)
            print(f"Documento modificado salvo em: {output_path}")
        except Exception as e:
            print(f"Erro ao salvar o documento final: {e}")


class SectionProcessor:
    """Responsável por manipular as propriedades de uma seção do documento."""

    def __init__(self, source_section, dest_section, nsmap: dict):
        self.source_sectPr = source_section._sectPr
        self.dest_sectPr = dest_section._sectPr
        self.nsmap = nsmap
        self.tags_to_copy = ['pgSz', 'pgMar', 'cols', 'docGrid', 'titlePg']

    def copy_properties(self):
        """Substitui as propriedades de layout da seção de destino pelas da origem."""
        print("  - Sincronizando propriedades da seção (margens, tamanho, etc.)...")

        # Remove propriedades antigas
        for child in list(self.dest_sectPr):
            if hasattr(child, 'tag') and etree.QName(child).localname in self.tags_to_copy:
                self.dest_sectPr.remove(child)

        # Adiciona novas propriedades
        for child in self.source_sectPr:
            if hasattr(child, 'tag') and etree.QName(child).localname in self.tags_to_copy:
                self.dest_sectPr.append(copy.deepcopy(child))

    def is_first_page_different(self) -> bool:
        """Verifica se a configuração 'Primeira Página Diferente' está ativa."""
        return self.source_sectPr.find('w:titlePg', namespaces=self.nsmap) is not None


class PartCopier:
    """Copia o conteúdo de uma parte do documento (cabeçalho/rodapé), lidando com a complexidade interna."""

    def __init__(self, source_element, dest_element, style_handler: 'StyleHandler', nsmap: dict, part_name: str):
        self.source_element = source_element
        self.dest_element = dest_element
        self.style_handler = style_handler
        self.nsmap = nsmap
        self.part_name = part_name

    def copy_content(self):
        """Executa o processo completo de cópia de conteúdo para esta parte."""
        if not self.source_element.paragraphs and not self.source_element.tables:
            print(f"  - A parte '{self.part_name}' está vazia no template. Pulando.")
            return

        print(f"  - Copiando conteúdo da parte: {self.part_name}...")

        rid_map = self._copy_image_relationships()
        dest_xml_element = self.dest_element._element
        dest_xml_element.clear()

        for child_element in self.source_element._element:
            processed_element = self._process_child_element(child_element, rid_map)
            dest_xml_element.append(processed_element)

    def _copy_image_relationships(self) -> dict:
        """Copia relações de imagem e retorna um mapa de IDs antigos para novos."""
        rid_map = {}
        source_part = self.source_element.part
        dest_part = self.dest_element.part
        for rel in source_part.rels.values():
            if "image" in rel.target_ref:
                rid_map[rel.rId] = dest_part.relate_to(rel.target_part, rel.reltype)
        return rid_map

    def _process_child_element(self, child_element, rid_map: dict):
        """Processa um único elemento filho (parágrafo, tabela), aplicando estilos e corrigindo referências."""
        # Cria uma cópia profunda para evitar modificar o original
        new_child_el = copy.deepcopy(child_element)

        # Se for um parágrafo, internaliza seus estilos
        if etree.QName(new_child_el).localname == 'p':
            self.style_handler.inline_paragraph_style(new_child_el)

        # Converte para string para corrigir IDs de imagem de forma segura
        child_xml_string = etree.tostring(new_child_el, encoding='unicode')
        for old_rid, new_rid in rid_map.items():
            child_xml_string = child_xml_string.replace(f'r:embed="{old_rid}"', f'r:embed="{new_rid}"')

        # Retorna um elemento lxml puro, pronto para ser anexado
        return etree.fromstring(child_xml_string)


class StyleHandler:
    """Classe utilitária para lidar com a complexa lógica de estilos do Word."""

    def __init__(self, template_doc, nsmap: dict):
        self.template_doc = template_doc
        self.nsmap = nsmap
        self.pPr_cache = {}  # Cache para propriedades de estilo de parágrafo

    def _fetch_style_pPr(self, style_id: str):
        """Busca (com cache) as propriedades de um estilo de parágrafo pelo seu ID."""
        if style_id in self.pPr_cache:
            return self.pPr_cache[style_id]

        try:
            style = self.template_doc.styles[style_id]
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                pPr = style.element.find('.//w:pPr', namespaces=self.nsmap)
                self.pPr_cache[style_id] = pPr
                return pPr
        except KeyError:
            pass  # Estilo não encontrado no documento

        self.pPr_cache[style_id] = None
        return None

    def inline_paragraph_style(self, p_element):
        """
        Copia as propriedades de um estilo referenciado (como paradas de tabulação)
        diretamente para o XML do parágrafo, tornando-o autossuficiente.
        """
        pPr = p_element.find('w:pPr', namespaces=self.nsmap)
        if pPr is None:
            pPr = etree.Element(f'{{{self.nsmap["w"]}}}pPr')
            p_element.insert(0, pPr)

        style_tag = pPr.find('w:pStyle', namespaces=self.nsmap)
        if style_tag is not None:
            style_id = style_tag.get(f'{{{self.nsmap["w"]}}}val')
            style_pPr = self._fetch_style_pPr(style_id)

            if style_pPr is not None:
                # Itera sobre as propriedades do estilo (ex: <w:tabs>)
                for style_prop in style_pPr:
                    prop_tag_name = etree.QName(style_prop).localname
                    # Se o parágrafo não tem essa propriedade, herda do estilo
                    if pPr.find(f'w:{prop_tag_name}', namespaces=self.nsmap) is None:
                        pPr.insert(0, copy.deepcopy(style_prop))


def synchronize_docx_layout(template_path: str, dest_path: str, output_path: str):
    """
    Função de alto nível que encapsula a criação e execução do sincronizador.
    Serve como uma API simples para o consumidor deste módulo.
    """
    if not all(os.path.exists(p) for p in [template_path, dest_path]):
        print("Erro: Arquivo de template ou destino não encontrado.")
        return

    try:
        synchronizer = DocxLayoutSynchronizer(template_path, dest_path)
        synchronizer.run()
        synchronizer.save(output_path)
    except Exception as e:
        print(f"Ocorreu um erro inesperado durante o processo: {e}")


# --- Função Principal ---
if __name__ == "__main__":
    TEMPLATE_FILE = "header1.docx"
    DEST_FILE = "original1.docx"
    OUTPUT_FILE = "resultado1.docx"

    # apagar arquivo destino
    if os.path.exists(OUTPUT_FILE):
        os.remove(OUTPUT_FILE)

    synchronize_docx_layout(TEMPLATE_FILE, DEST_FILE, OUTPUT_FILE)